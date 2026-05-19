"""
Resume text extraction utilities.
Supports PDF (via pdfminer or basic extraction), DOCX, and plain text.
"""

import re
import os
import io


def extract_text_from_txt(file_bytes):
    """Extract text from a plain text file."""
    try:
        return file_bytes.decode('utf-8', errors='ignore')
    except Exception:
        return ""


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF using available libraries."""
    text = ""
    
    # Try pdfminer.six first (best quality)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        from pdfminer.layout import LAParams
        text = pdfminer_extract(io.BytesIO(file_bytes), laparams=LAParams())
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Try PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text += page.extract_text() or ""
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Try pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text += page.extract_text() or ""
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: basic text extraction from PDF bytes
    try:
        raw = file_bytes.decode('latin-1', errors='ignore')
        # Extract text between stream/endstream markers
        streams = re.findall(r'stream\r?\n(.*?)\r?\nendstream', raw, re.DOTALL)
        extracted = []
        for stream in streams:
            # Get printable ASCII text
            printable = re.sub(r'[^\x20-\x7E\n]', ' ', stream)
            words = re.findall(r'[A-Za-z][A-Za-z0-9\s,.\-@]+', printable)
            if words:
                extracted.append(' '.join(words))
        text = ' '.join(extracted)
    except Exception:
        pass
    
    return text


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also get table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return '\n'.join(paragraphs)
    except ImportError:
        pass
    except Exception:
        pass
    
    # Fallback: extract from XML
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            with zf.open('word/document.xml') as doc_xml:
                tree = ET.parse(doc_xml)
                root = tree.getroot()
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                texts = []
                for elem in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if elem.text:
                        texts.append(elem.text)
                return ' '.join(texts)
    except Exception:
        pass
    
    return ""


def extract_resume_text(file_bytes, filename):
    """
    Main entry point: extract text from a resume file.
    Supports PDF, DOCX, DOC, TXT.
    """
    ext = os.path.splitext(filename.lower())[1]
    
    if ext == '.pdf':
        text = extract_text_from_pdf(file_bytes)
    elif ext in ('.docx', '.doc'):
        text = extract_text_from_docx(file_bytes)
    elif ext in ('.txt', '.text'):
        text = extract_text_from_txt(file_bytes)
    else:
        # Try as text
        text = extract_text_from_txt(file_bytes)
    
    return clean_text(text)


def clean_text(text):
    """Clean and normalize extracted text."""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    # Remove very long strings of repeated characters (artifacts)
    text = re.sub(r'(.)\1{5,}', r'\1', text)
    
    return text.strip()
